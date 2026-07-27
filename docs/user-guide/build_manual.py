from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "대방동지역주택조합_ERP_사용설명서_초안.docx"
MODULES = [
    "01-quick-start.md",
    "02-basic-info.md",
    "03-approval.md",
    "04-finance.md",
    "05-members.md",
    "06-hr-reports.md",
    "07-roles-status.md",
    "08-troubleshooting.md",
]

BLUE = "2457C5"
DARK = "152238"
MUTED = "667085"
PALE_BLUE = "EAF2FF"
PALE_GRAY = "F4F6F9"
BORDER = "D9E2EC"


def set_font(run, size=None, bold=None, color=None, name="Malgun Gothic"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(widths)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])


def create_numbering(doc, bullet=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, indent])
    level.append(p_pr)
    abstract.append(level)
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def add_numbered_paragraph(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.18
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    paragraph._p.get_or_add_pPr().append(num_pr)
    add_inline_runs(paragraph, text)
    return paragraph


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, color, before, after in (
        ("Heading 1", 18, BLUE, 16, 8),
        ("Heading 2", 14, BLUE, 12, 6),
        ("Heading 3", 11.5, DARK, 9, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("대방동지역주택조합 ERP | 사용자 튜토리얼"), 8.5, False, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("사용설명서 초안  ·  "), 8.5, False, MUTED)
    add_page_field(footer)


def add_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(kicker.add_run("USER GUIDE · TUTORIAL EDITION"), 10, True, BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(10)
    set_font(title.add_run("대방동지역주택조합 ERP"), 27, True, DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    set_font(subtitle.add_run("처음부터 업무 완료까지 따라 하는 쉬운 사용설명서"), 14, False, BLUE)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_after = Pt(18)
    set_font(note.add_run("현재 운영 기능 기준 · 지속 업데이트형 초안"), 10.5, True, MUTED)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [2100, 6900])
    rows = [
        ("운영 주소", "https://dbapt-erp.vercel.app/"),
        ("적용 기준", "2026년 7월 17일 운영 버전"),
        ("문서 관리", "docs/user-guide의 기능별 Markdown 원본에서 갱신"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        keep_row_together(row)
        set_cell_fill(row.cells[0], PALE_BLUE)
        set_font(row.cells[0].paragraphs[0].add_run(label), 9.5, True, BLUE)
        set_font(row.cells[1].paragraphs[0].add_run(value), 9.5, False, DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("이 설명서는 화면 위치보다 실제 업무 흐름과 완료 기준을 중심으로 설명합니다."), 9.5, False, MUTED)
    doc.add_page_break()


def add_toc(doc):
    doc.add_heading("사용 순서", level=1)
    items = [
        "1. 빠른 시작과 화면 구성",
        "2. 기초정보 준비",
        "3. 기안·결재",
        "4. 회계·지급·전표",
        "5. 조합원 관리",
        "6. 인사·급여와 보고서",
        "7. 역할별 업무와 상태",
        "8. 문제 해결",
    ]
    toc_num_id = create_numbering(doc)
    for item in items:
        add_numbered_paragraph(doc, re.sub(r"^\d+\.\s*", "", item), toc_num_id)

    doc.add_heading("핵심 업무 흐름", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    set_table_widths(table, [1500, 1500, 1500, 1500, 1500, 1500])
    labels = ["기안 작성", "전자결재", "회의 의결", "지출결의", "실제 지급", "전표 확인"]
    for index, (cell, label) in enumerate(zip(table.rows[0].cells, labels)):
        set_cell_fill(cell, PALE_BLUE if index % 2 == 0 else PALE_GRAY)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(cell.paragraphs[0].add_run(label), 8.5, True, BLUE if index % 2 == 0 else DARK)

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_table_widths(callout, [9000])
    keep_row_together(callout.rows[0])
    set_cell_fill(callout.cell(0, 0), "FFF7DF")
    p = callout.cell(0, 0).paragraphs[0]
    set_font(p.add_run("중요  "), 10, True, "8A5A00")
    set_font(p.add_run("결재 승인, 회의 의결, 지출결의 승인, 실제 지급과 전표 생성은 서로 다른 단계입니다."), 10, False, DARK)
    doc.add_page_break()


def add_inline_runs(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, 9.5, True, BLUE)
            run.font.highlight_color = None
        elif part.startswith("**") and part.endswith("**"):
            set_font(paragraph.add_run(part[2:-2]), 10.5, True, DARK)
        else:
            set_font(paragraph.add_run(part), 10.5, False, DARK)


def append_markdown(doc, path):
    lines = path.read_text(encoding="utf-8").splitlines()
    numbered_id = None
    bullet_id = None
    previous_was_number = False
    previous_was_bullet = False
    for line in lines:
        text = line.strip()
        if not text:
            previous_was_number = False
            previous_was_bullet = False
            continue
        if text.startswith("# "):
            doc.add_heading(text[2:], level=1)
        elif text.startswith("## "):
            doc.add_heading(text[3:], level=2)
        elif text.startswith("### "):
            doc.add_heading(text[4:], level=3)
        elif re.match(r"^\d+\.\s+", text):
            if not previous_was_number:
                numbered_id = create_numbering(doc)
            add_numbered_paragraph(doc, re.sub(r"^\d+\.\s+", "", text), numbered_id)
            previous_was_number = True
            continue
        elif text.startswith("- "):
            if not previous_was_bullet:
                bullet_id = create_numbering(doc, bullet=True)
            add_numbered_paragraph(doc, text[2:], bullet_id)
            previous_was_bullet = True
            previous_was_number = False
            continue
        elif text.startswith("주의:") or text.startswith("중요:") or text.startswith("완료 기준:"):
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            set_table_widths(table, [9000])
            keep_row_together(table.rows[0])
            set_cell_fill(table.cell(0, 0), "FFF7DF" if text.startswith(("주의:", "중요:")) else PALE_BLUE)
            p = table.cell(0, 0).paragraphs[0]
            label, value = text.split(":", 1)
            set_font(p.add_run(f"{label}  "), 10, True, "8A5A00" if label in ("주의", "중요") else BLUE)
            add_inline_runs(p, value.strip())
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, text)
        previous_was_number = False
        previous_was_bullet = False


def add_maintenance_appendix(doc):
    doc.add_heading("부록. 설명서 업데이트 방법", level=1)
    paragraphs = [
        "기능을 추가하거나 메뉴명이 바뀌면 관련 Markdown 파일만 수정하고 CHANGELOG.md에 변경 내용을 기록합니다.",
        "사이드바 전체 이미지는 공통 문서에 한 번만 사용하고, 각 튜토리얼에서는 ‘기안·결재 → 새 기안’처럼 텍스트 경로를 사용합니다.",
        "정식 배포 전에는 운영 화면에서 모든 경로와 버튼명을 다시 확인하고, 개인정보를 가린 최종 스크린샷으로 교체합니다.",
        "Word 설명서는 이 폴더의 build_manual.py를 실행해 다시 생성한 뒤 전 페이지 렌더링 검수를 진행합니다.",
    ]
    appendix_num_id = create_numbering(doc)
    for item in paragraphs:
        add_numbered_paragraph(doc, item, appendix_num_id)

    doc.add_heading("이미지 촬영 체크리스트", level=2)
    checklist_bullet_id = create_numbering(doc, bullet=True)
    for item in (
        "실제 조합원 이름, 주민번호, 연락처와 계좌번호를 가렸는가?",
        "화면의 클릭 위치와 설명의 단계 번호가 일치하는가?",
        "현재 운영 버전의 메뉴명과 선택 상태가 정확한가?",
        "같은 사이드바 이미지를 여러 장에서 중복 사용하지 않았는가?",
        "이미지 아래에 그림 설명과 대체 텍스트가 있는가?",
    ):
        add_numbered_paragraph(doc, item, checklist_bullet_id)


def build():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "대방동지역주택조합 ERP 사용설명서"
    doc.core_properties.subject = "튜토리얼형 지속 업데이트 사용설명서"
    doc.core_properties.author = "대방동지역주택조합"
    add_cover(doc)
    add_toc(doc)
    for module in MODULES:
        append_markdown(doc, ROOT / module)
    add_maintenance_appendix(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
